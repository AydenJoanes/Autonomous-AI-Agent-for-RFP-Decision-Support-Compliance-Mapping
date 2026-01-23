"""
Document Parser Factory - Unified interface for parsing documents
"""

from pathlib import Path
from loguru import logger
import re

class DocumentParserFactory:
    """
    Factory for creating document parsers and handling file parsing.
    Supports PDF and DOCX.
    """
    
    @staticmethod
    def parse_with_fallback(file_path: str) -> str:
        """
        Parse a file and return its text content.
        Automatically detects file type.
        
        Args:
            file_path: Absolute path to the file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
            
        suffix = path.suffix.lower()
        
        try:
            if suffix == ".pdf":
                return DocumentParserFactory._parse_pdf(path)
            elif suffix in [".docx", ".doc"]:
                return DocumentParserFactory._parse_docx(path)
            elif suffix in [".txt", ".md"]:
                return path.read_text(encoding="utf-8")
            else:
                logger.warning(f"Unsupported file format: {suffix}, attempting text read")
                return path.read_text(encoding="utf-8", errors="ignore")
                
        except Exception as e:
            logger.error(f"Parsing failed for {file_path}: {e}")
            raise

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Parse PDF using pypdf."""
        logger.info(f"Parsing PDF: {path.name}")
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            logger.info(f"Starting PDF extraction for: {path.name} with {len(reader.pages)} pages")
            text = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                        logger.debug(f"Page {i+1}/{len(reader.pages)} extracted {len(page_text)} chars")
                    else:
                        logger.warning(f"Page {i+1}/{len(reader.pages)} text extraction returned empty")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1}: {e}. Skipping page.")
            
            full_text = "\n\n".join(text)
            logger.info(f"PDF Extraction Complete. Total chars: {len(full_text)}. Pages attempted: {len(reader.pages)}")
            return full_text
        except ImportError:
            logger.error("pypdf not installed. Run: pip install pypdf")
            raise

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Parse DOCX using python-docx."""
        logger.info(f"Parsing DOCX: {path.name}")
        try:
            from docx import Document
            doc = Document(str(path))
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        text.append(" | ".join(row_data))
                        
            full_text = "\n\n".join(text)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX")
            return full_text
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()
