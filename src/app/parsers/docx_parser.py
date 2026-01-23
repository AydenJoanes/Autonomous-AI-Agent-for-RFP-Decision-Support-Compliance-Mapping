from src.app.parsers.base_parser import BaseParser
import docx
from loguru import logger
import os

class DocxParser(BaseParser):
    
    def supports_format(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in ['.docx', '.doc']

    def parse(self, file_path: str) -> str:
        logger.info(f"Using python-docx fallback parser for {file_path}")
        self.validate_file(file_path)
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract tables and convert to simple markdown tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    # Simple pipe separated table row
                    full_text.append("| " + " | ".join(row_text) + " |")
                full_text.append("") # space after table

            return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Docx parsing failed for {file_path}: {e}")
            raise RuntimeError(f"Docx failed to parse {file_path}: {str(e)}") from e
